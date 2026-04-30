"""
Text extraction utilities for PDF and DOCX files.
"""

import logging
import io
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from PDF bytes using PyMuPDF (fitz).
    Preserves all content including headers, bullets, tables.
    """
    try:
        import fitz  # PyMuPDF
        
        text_parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text with layout preservation
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---")
                text_parts.append(text)
        
        doc.close()
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            logger.warning("PDF extraction returned empty text")
            return ""
        
        logger.info(f"PDF extracted: {len(full_text)} characters from {len(text_parts)} pages")
        return full_text
        
    except ImportError:
        logger.error("PyMuPDF (fitz) not installed")
        raise RuntimeError("PDF extraction library not available")
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Could not extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from DOCX bytes using python-docx.
    Includes paragraphs and table cells.
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract table content (important for resume tables)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        
        full_text = "\n".join(text_parts)
        
        if not full_text.strip():
            logger.warning("DOCX extraction returned empty text")
            return ""
        
        logger.info(f"DOCX extracted: {len(full_text)} characters")
        return full_text
        
    except ImportError:
        logger.error("python-docx not installed")
        raise RuntimeError("DOCX extraction library not available")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"Could not extract text from DOCX: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Main entry point. Routes to correct extractor based on file type.
    file_type: 'pdf' or 'docx'
    """
    file_type = file_type.lower().strip(".")
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Only PDF and DOCX supported.")


def clean_extracted_text(raw_text: str) -> str:
    """
    Clean extracted text while preserving all information.
    Removes excessive whitespace but keeps structure.
    """
    if not raw_text:
        return ""
    
    lines = raw_text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        cleaned = line.strip()
        cleaned_lines.append(cleaned)
    
    # Remove more than 2 consecutive empty lines
    result_lines = []
    empty_count = 0
    for line in cleaned_lines:
        if not line:
            empty_count += 1
            if empty_count <= 2:
                result_lines.append(line)
        else:
            empty_count = 0
            result_lines.append(line)
    
    return "\n".join(result_lines)
